# -*- coding: utf-8 -*-
"""
build_docx_v3.py — 通用附件7分析报告生成器 v3
> 操作时间: 2026-06-30 15:00
> 操作agent: default

从 config.json + 达成度 xlsx → 生成附件7 docx

用法：python build_docx_v3.py config.json template.docx xlsx_path output.docx

config.json 需包含: course_name, course_code, term, assessment, major, classes[], 
   teacher, hours, credits, kaohe_desc, targets[], students[]
"""
import json, openpyxl, sys, os, math, copy
from docx import Document
from docx.shared import Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH

NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def fill_label_cell(cell, append_val):
    """追加内容到单元格，保留模板原有标签"""
    text = cell.text.strip()
    if str(append_val) in text:
        return
    para = cell.paragraphs[0]
    if not para.runs:
        para.add_run(str(append_val))
    elif para.runs[-1].text.strip() == '':
        para.runs[-1].text = str(append_val)
    else:
        para.add_run(str(append_val))

def remove_vmerge_from_row(row):
    """移除行内所有单元格的 vMerge，避免孤立的合并标记"""
    for cell in row.cells:
        tc_pr = cell._tc.find(f'{{{NS_W}}}tcPr')
        if tc_pr is not None:
            for vm in tc_pr.findall(f'{{{NS_W}}}vMerge'):
                tc_pr.remove(vm)

def set_cell_text_clean(cell, text, font_name='Times New Roman', font_size=10.5):
    for p in cell.paragraphs:
        p.clear()
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(font_size)
    run.font.name = font_name

def fill_stat_row(cell, max_v, min_v, avg_v, pass_rate):
    """填充统计行：保留模板标签，替换占位为实际值"""
    para = cell.paragraphs[0]
    para.clear()
    text = f"最高分：{int(max_v)}    最低分：{int(min_v)}    平均分：{avg_v:.1f}    及格率：{pass_rate:.1f}%"
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10.5)

def append_analysis_para(cell, text):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.first_line_indent = Emu(327660)
    pf.line_spacing = 1.5
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '仿宋'
    return p

def get_distribution(scores):
    n = len(scores)
    if n == 0:
        return {'total': 0, 'max': 0, 'min': 0, 'avg': 0, 'pass_rate': 0, 'counts': [0]*7, 'pcts': [0]*7}
    max_s, min_s = max(scores), min(scores)
    avg_s = sum(scores) / n
    pass_rate = len([s for s in scores if s >= 60]) / n * 100
    counts = [0] * 7
    for s in scores:
        if s >= 90: counts[0] += 1
        elif s >= 80: counts[1] += 1
        elif s >= 70: counts[2] += 1
        elif s >= 60: counts[3] += 1
        elif s >= 50: counts[4] += 1
        elif s >= 40: counts[5] += 1
        else: counts[6] += 1
    pcts = [c / n * 100 for c in counts]
    return {'total': n, 'max': max_s, 'min': min_s, 'avg': avg_s, 'pass_rate': pass_rate, 'counts': counts, 'pcts': pcts}

def build_docx_v3(config_path, template_path, xlsx_path, output_path):
    """主函数：生成附件7"""
    with open(config_path, 'r', encoding='utf-8') as f:
        cfg = json.load(f)
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    ws = wb.active

    tg = cfg['targets']
    nt = len(tg)
    students = cfg['students']
    n_all = len(students)
    all_scores = [st['luru_qm'] for st in students]  # 终结性考核用期末成绩
    d_all = get_distribution(all_scores)

    # 分班
    class_students = {}
    for st in students:
        cls = st.get('class', cfg['classes'][0])
        class_students.setdefault(cls, []).append(st)

    # 分班分布
    class_dists = {}
    for cls, sts in class_students.items():
        class_dists[cls] = get_distribution([st['luru_qm'] for st in sts])

    # 每条考核环节的折合分列数（每目标5项）
    n_cols_p_t = [len(t['items']) for t in tg]

    # 各考核环节平均折合分
    all_item_avgs = {}
    for ti, t in enumerate(tg):
        for ii, it in enumerate(t['items']):
            vals = []
            for st in students:
                raw = st['raw_scores'].get(it['name'], 0)
                zh = raw / it['raw_full'] * it['zhe_full'] if it['raw_full'] else 0
                vals.append(zh)
            all_item_avgs[(ti, ii)] = round(sum(vals) / len(vals), 2) if vals else 0

    # 教学目标达成度
    target_dcs = {}
    for ti, t in enumerate(tg):
        sum_avg = sum(all_item_avgs[(ti, ii)] for ii in range(n_cols_p_t[ti]))
        sum_full = sum(it['zhe_full'] for it in t['items'])
        target_dcs[t['name']] = round(sum_avg / sum_full, 2) if sum_full > 0 else 0

    # 课程达成度
    course_dc = round(sum(target_dcs.values()) / len(target_dcs), 2) if target_dcs else 0

    # 分班达成度
    class_dcs = {}
    for cls_name in sorted(class_students.keys()):
        cls_sts = class_students[cls_name]
        cls_target_dcs = []
        for ti, t in enumerate(tg):
            sum_avg = 0
            for ii, it in enumerate(t['items']):
                vals = []
                for st in cls_sts:
                    raw = st['raw_scores'].get(it['name'], 0)
                    zh = raw / it['raw_full'] * it['zhe_full'] if it['raw_full'] else 0
                    vals.append(zh)
                sum_avg += sum(vals) / len(vals) if vals else 0
            sum_full = sum(it['zhe_full'] for it in t['items'])
            dc = round(sum_avg / sum_full, 2) if sum_full > 0 else 0
            cls_target_dcs.append(dc)
        cls_cd = round(sum(cls_target_dcs) / len(cls_target_dcs), 2) if cls_target_dcs else 0
        class_dcs[cls_name] = {'targets': cls_target_dcs, 'course': cls_cd}

    # ═══════════ 打开模板 ═══════════
    doc = Document(template_path)
    t0 = doc.tables[0]  # 成绩分布表
    t1 = doc.tables[1]  # 达成度表

    # ═══════════ Table0: 成绩分布（匹配模板真实行列）═══════════
    # 列映射：7 个分数段的列索引
    col_keys = [2, 3, 5, 6, 7, 9, 10]  # 90-100, 80-89, 70-79, 60-69, 50-59, 40-49, <40
    
    # R0: 课程名称(C0-C4) | 课程代码(C5-C7) | 课程性质(C8-C10)
    fill_label_cell(t0.rows[0].cells[0], cfg.get('course_name', ''))
    fill_label_cell(t0.rows[0].cells[5], cfg.get('course_code', ''))
    fill_label_cell(t0.rows[0].cells[8], '必修课（通识教育课）')
    
    # R1: 开课部门(C0-C4) | 考核方式(C5-C7) | 学分(C8-C10)
    fill_label_cell(t0.rows[1].cells[0], cfg.get('major', ''))
    fill_label_cell(t0.rows[1].cells[5], cfg.get('assessment', ''))
    fill_label_cell(t0.rows[1].cells[8], cfg.get('credits', ''))
    
    # R2: 教学班(C1-C5) | 应考人数(C6-C8) | 实考人数(C9-C10)
    all_classes_str = '、'.join(c + '班' for c in cfg.get('classes', []))
    fill_label_cell(t0.rows[2].cells[1], all_classes_str)
    fill_label_cell(t0.rows[2].cells[6], str(n_all))
    fill_label_cell(t0.rows[2].cells[9], str(n_all))
    
    # R4: 人数, R5: 百分比, R6: 统计
    for ci, cnt in enumerate(d_all['counts']):
        fill_label_cell(t0.rows[4].cells[col_keys[ci]], str(cnt))
        fill_label_cell(t0.rows[5].cells[col_keys[ci]], f"{d_all['pcts'][ci]:.1f}%")
    fill_stat_row(t0.rows[6].cells[1], d_all['max'], d_all['min'], d_all['avg'], d_all['pass_rate'])
    
    # ══ 行政班成绩分布 ══
    class_names = list(class_students.keys())
    cls_base_rows = [(8, 10, 11, 12), (13, 15, 16, 17)]  # (标题行, 人数行, 百分比行, 统计行)
    
    for ci in range(min(len(class_names), 2)):
        cname = class_names[ci]
        d = class_dists[cname]
        n = d['total']
        title_row, cnt_row, pct_row, stat_row = cls_base_rows[ci]
        
        # 行政班(C1-C3) | 任课教师(C4-C5) | 应考(C6-C8) | 实考(C9-C10)
        fill_label_cell(t0.rows[title_row].cells[1], cname + '班')
        fill_label_cell(t0.rows[title_row].cells[4], cfg.get('teacher', ''))
        fill_label_cell(t0.rows[title_row].cells[6], str(n))
        fill_label_cell(t0.rows[title_row].cells[9], str(n))
        
        for cj, cnt in enumerate(d['counts']):
            fill_label_cell(t0.rows[cnt_row].cells[col_keys[cj]], str(cnt))
            fill_label_cell(t0.rows[pct_row].cells[col_keys[cj]], f"{d['pcts'][cj]:.1f}%")
        fill_stat_row(t0.rows[stat_row].cells[1], d['max'], d['min'], d['avg'], d['pass_rate'])

    # ═══════════ Table1: 达成度 ═══════════
    # 1. 找到摘要行（".. ..." 占位符）
    summary_row_idx = None
    for ri, row in enumerate(t1.rows):
        c1 = row.cells[1].text.strip() if len(row.cells) > 1 else ''
        if '..' in c1 and len(c1) <= 10 and ri > 3:
            summary_row_idx = ri
            break
    if summary_row_idx is None:
        summary_row_idx = 1 + sum(len(t['items']) for t in tg)
    
    # 2. 删掉模板的旧数据行（R1 ~ R{summary_row_idx-1}）
    trs_to_remove = [t1.rows[i]._tr for i in range(1, summary_row_idx)]
    for tr in trs_to_remove:
        t1._tbl.remove(tr)
    summary_row_idx = 1  # 摘要行现在在 R1 位置
    # 移除表头 vMerge，避免孤立的合并标记
    remove_vmerge_from_row(t1.rows[0])
    
    # 3. 用 add_row() 创建全部数据行（加到表尾，再移到摘要前行）
    total_rows = sum(len(t['items']) for t in tg)
    for _ in range(total_rows):
        t1.add_row()
    
    # 移到摘要前行：取最后 total_rows 个 tr
    all_trs = t1._tbl.findall(f'{{{NS_W}}}tr')
    data_trs = all_trs[-total_rows:]
    for tr in data_trs:
        t1._tbl.remove(tr)
        summary_tr = t1.rows[summary_row_idx]._tr  # R1 = 摘要行
        summary_tr.addprevious(tr)
        summary_row_idx += 1
    
    # 4. 填充数据行（R1~R{total_rows}）
    ri_data = 1
    for ti, t in enumerate(tg):
        for ii, it in enumerate(t['items']):
            row = t1.rows[ri_data]
            
            # C0: 考核环节（仅第一行）
            if ri_data == 1:
                set_cell_text_clean(row.cells[0], '考核环节')
            
            # C1-C2: 毕业要求+指标点（每目标首行）
            if ii == 0:
                set_cell_text_clean(row.cells[1], '毕业要求1')
                set_cell_text_clean(row.cells[2], t.get('indicator', ''))
            
            # C3-C4: 教学目标名（每目标首行）
            if ii == 0:
                set_cell_text_clean(row.cells[3], t['name'])
                set_cell_text_clean(row.cells[4], t['name'])
            
            # C5-C7: 考核环节名称 + 原始满分 + 折合满分
            set_cell_text_clean(row.cells[5], it['name'])
            set_cell_text_clean(row.cells[6], str(it['raw_full']))
            set_cell_text_clean(row.cells[7], str(it['zhe_full']))
            
            # C8: 平均分
            avg_val = all_item_avgs.get((ti, ii), 0)
            set_cell_text_clean(row.cells[8], str(avg_val))
            
            # C9: 达成度
            dc_val = avg_val / it['zhe_full'] if it['zhe_full'] else 0
            set_cell_text_clean(row.cells[9], f"{dc_val:.1%}")
            
            ri_data += 1
    
    # 5. 填充摘要行（summary_row_idx 已在上面追踪到正确位置）
    set_cell_text_clean(t1.rows[summary_row_idx].cells[1], '课程教学目标达成度')
    set_cell_text_clean(t1.rows[summary_row_idx].cells[9], f'{course_dc:.2f}')
    
    analysis_row_idx = summary_row_idx + 1
    
    # ═══════════ 补充分析计算 ═══════════
    # 各考核环节达成率
    item_rates = {}
    for ti, t in enumerate(tg):
        for ii, it in enumerate(t['items']):
            avg = all_item_avgs.get((ti, ii), 0)
            rate = avg / it['zhe_full'] if it['zhe_full'] else 0
            item_rates[(ti, ii)] = {
                'tname': t['name'], 'iname': it['name'],
                'avg': avg, 'full': it['zhe_full'], 'rate': rate
            }
    
    # 最薄弱的三项（按达成率升序）
    weakest_items = sorted(item_rates.values(), key=lambda x: x['rate'])[:3]
    
    # 各目标最弱考核环节
    target_weakest = {}
    for ti, t in enumerate(tg):
        items_t = [(ii, item_rates[(ti, ii)]) for ii in range(len(t['items']))]
        items_t.sort(key=lambda x: x[1]['rate'])
        target_weakest[t['name']] = items_t[0][1]  # 最弱的一项
    
    # 教学目标排名
    target_ranking = sorted(target_dcs.items(), key=lambda x: x[1])
    
    # 低于60分和低于50分
    below_60 = 100 - d_all['pass_rate']
    below_50 = d_all['pcts'][5] + d_all['pcts'][6]
    
    # 低分学生统计（期末成绩<60）
    low_students = [st for st in students if st.get('luru_qm', 0) < 60]
    low_student_names = '、'.join(st['name'] for st in low_students[:5])
    if len(low_students) > 5:
        low_student_names += f'等{len(low_students)}人'
    
    # 分班低分对比（期末成绩）
    class_low = {}
    for cls, sts in class_students.items():
        low_in_cls = [st for st in sts if st.get('luru_qm', 0) < 60]
        class_low[cls] = len(low_in_cls)
    
    # ═══════════ 分析文本 ═══════════
    classes_desc = '、'.join(c + '班' for c in cfg.get('classes', []))
    targets_dc_text = '，'.join([f"{t['name']}={target_dcs.get(t['name'], 0):.2f}" for t in tg])
    
    # 分达成度档次描述
    dc_level = '较高' if course_dc >= 0.75 else ('中等' if course_dc >= 0.65 else '偏低')
    
    # 分班对比文本
    class_comp = ''
    if len(class_dcs) >= 2:
        cls_names = sorted(class_dcs.keys())
        c1, c2 = cls_names[0], cls_names[1]
        dc1 = class_dcs[c1]['course']
        dc2 = class_dcs[c2]['course']
        diff = abs(dc1 - dc2)
        class_comp = f"\n\n分班达成度对比：{c1}班={dc1:.2f}，{c2}班={dc2:.2f}，差异为{diff:.2f}。"
        if diff > 0.05:
            class_comp += f" {c2}班达成度明显低于{c1}班（差距>{0.05:.0%}），需重点关注该班教学过程和学习效果。"
        # 分班低分详情
        if class_low:
            low_parts = [f"{c}班{cn}人" for c, cn in class_low.items() if cn > 0]
            if low_parts:
                class_comp += f" 不及格学生分布：{'，'.join(low_parts)}。"
    
    # 分班低分简述（单班情况）
    class_low_text = ''
    if len(class_dcs) < 2 and class_low:
        for cls, cn in class_low.items():
            if cn > 0:
                class_low_text = f"其中{cls}班不及格{cn}人（{'、'.join(st['name'] for st in low_students[:3])}等），"
    
    analysis_text = f"""一、终结性考核典型性错误及试题难度分析

（一）终结性考核概况
本课程终结性考核采用期末考试形式（闭卷笔试，满分100分），面向{classes_desc}共{n_all}名学生。考试最高分{int(d_all['max'])}分，最低分{int(d_all['min'])}分，平均分{d_all['avg']:.1f}分，及格率{d_all['pass_rate']:.1f}%。从整体成绩分布来看，{class_low_text}考试成绩呈现"{'两极分化较明显' if int(d_all['max'])-int(d_all['min']) >= 40 else '分布相对集中'}"的特征。

（二）成绩分段分布
90-100分{d_all['counts'][0]}人（{d_all['pcts'][0]:.1f}%），80-89分{d_all['counts'][1]}人（{d_all['pcts'][1]:.1f}%），70-79分{d_all['counts'][2]}人（{d_all['pcts'][2]:.1f}%），60-69分{d_all['counts'][3]}人（{d_all['pcts'][3]:.1f}%），50-59分{d_all['counts'][4]}人（{d_all['pcts'][4]:.1f}%），40-49分{d_all['counts'][5]}人（{d_all['pcts'][5]:.1f}%），40分以下{d_all['counts'][6]}人（{d_all['pcts'][6]:.1f}%）。{'及格率超过85%，整体教学效果良好。' if d_all['pass_rate'] >= 85 else ('及格率处于中等水平，存在一定提升空间。' if d_all['pass_rate'] >= 70 else '不及格比例偏高，需引起高度重视。')}

（三）典型性错误与试题难度分析
从考试结果看，{below_60:.0f}%的学生成绩低于60分{'，低于50分的有' + str(int(below_50)) + '%' if below_50 > 0 else ''}。不及格学生{'主要集中在' + low_student_names + '，' if low_students else ''}反映出部分学生在基础概念理解、计算准确性、综合应用能力等方面存在明显不足。结合各教学目标对应的考核子项分析，期末考试作为覆盖全部教学目标的终结性评价，其得分率直接影响了各目标的达成度水平。建议教师在阅卷过程中对典型性错误进行分类记录，分析错误类型（概念性错误、计算性错误、审题偏差等），以便在后续教学中开展针对性纠错训练。

二、课程教学目标达成分析

本课程共设{nt}个教学目标，对应课程教学质量标准中规定的毕业要求指标点。各教学目标达成度如下：{targets_dc_text}。课程整体达成度为{course_dc:.2f}，处于{dc_level}水平，{'整体教学目标得以较好实现' if course_dc >= 0.75 else ('教学目标基本达成，部分目标需加强' if course_dc >= 0.65 else '部分教学目标达成度偏低，亟待改进')}。

在各教学目标中，达成度最高的是{target_ranking[-1][0]}（{target_ranking[-1][1]:.2f}），达成度最低的是{target_ranking[0][0]}（{target_ranking[0][1]:.2f}）。造成目标间差异的主要原因是各目标对应的考核环节得分率不同：{target_ranking[0][0]}中得分率最低的考核环节为"{target_weakest[target_ranking[0][0]]['iname']}"（达成率{target_weakest[target_ranking[0][0]]['rate']:.1%}），拉低了该目标的整体达成度；而{target_ranking[-1][0]}中各考核环节表现相对均衡。

从考核环节维度分析，得分率最低的三个环节为：{'；'.join(f'"{w["tname"]}-{w["iname"]}"达成率{w["rate"]:.1%}（均分{w["avg"]:.1f}/{w["full"]:.1f}）' for w in weakest_items)}。其中"{weakest_items[0]['iname']}"作为{weakest_items[0]['tname']}的关键考核项，表现尤为薄弱，建议作为课程改进的重中之重。{class_comp}

三、学生学习和教师教学过程中存在的问题

（一）学生学习方面
1. 基础知识掌握不牢固。部分学生在{'阶段测试' if '阶段测试' in weakest_items[0]['iname'] else weakest_items[0]['iname']}中得分偏低（达成率仅{weakest_items[0]['rate']:.1%}），反映出基础概念、基本技能掌握不扎实。
2. 综合应用能力有待提升。{f'期末考试成绩分布跨度大（{int(d_all["max"])}分-{int(d_all["min"])}分），' if int(d_all['max'])-int(d_all['min']) >= 30 else ''}说明学生综合运用知识解决实际问题的能力参差不齐。
3. 自主学习与持续改进意识不足。{f'不及格学生{len(low_students)}人（占{len(low_students)/n_all*100:.1f}%），' if low_students else ''}表明部分学生未能在课程学习过程中及时巩固知识、弥补薄弱环节。
4. {'平时成绩与期末成绩的关联性值得关注，部分学生平时考核表现较好但期末考试成绩不理想，可能存在临时抱佛脚或平时考核未能真实反映学习效果的情况。' if d_all['pass_rate'] < 90 else ''}

（二）教师教学方面
1. 针对薄弱知识点的专项训练不足。达成度最低的考核环节集中在{'、'.join(w['iname'] for w in weakest_items[:2])}等方面，说明教学中对这些环节的强化不够。
2. 过程性关注和帮扶机制有待加强。对于学情预警的响应不够及时，部分学生从学期初就表现出学习困难但未得到有效干预。
3. 教学资源布局有待优化。可进一步丰富视频讲解、在线练习、错题分析等辅导资源，特别是针对薄弱环节的补充材料。
4. 不同班级间教学效果的均衡性需关注（如达成度差异、不及格率差异），应统一教学要求并加强薄弱班级的教学督导。

四、对今后教学的改进意见

1. 强化薄弱环节专项训练。针对{weakest_items[0]['tname']}的"{weakest_items[0]['iname']}"（达成率仅{weakest_items[0]['rate']:.1%}）和{'、'.join(f'"{w["iname"]}"（{w["rate"]:.1%}）' for w in weakest_items[1:2])}，设计专项练习和阶段性测验，并纳入平时成绩考核。
2. 建立学业预警与帮扶机制。定期（每4周）分析学生成绩数据，对{'折合得分偏低或单项落后' if not low_students else f'不及格风险学生（如{low_student_names}）'}进行重点关注，安排课后辅导或学习小组，确保及时发现、及时干预。
3. 优化教学资源配置。录制薄弱知识点精讲视频，建设在线自测题库，提供错题解析和针对性练习，为自主学习提供支撑。
4. 加强过程性考核的诊断功能。在学期中增加一次模拟测验，让学生了解自身薄弱点；对阶段测试结果进行知识点归因分析，调整后续教学重点。
5. 推进班级教学均衡。{'统一教学进度和考核标准，对教学效果偏低的班级增加课堂互动和随堂测验频次。' if len(class_dcs) >= 2 else '持续关注全班学生学习动态，对后进生实行"一对一"帮扶。'}"""

    # 写入分析区
    if analysis_row_idx < len(t1.rows) and len(t1.rows[analysis_row_idx].cells) >= 2:
        cell = t1.rows[analysis_row_idx].cells[1]
        for p in cell.paragraphs:
            p.clear()
        sections = analysis_text.split('\n\n')
        for i, sec in enumerate(sections):
            if i == 0:
                p = cell.paragraphs[0]
            else:
                p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf = p.paragraph_format
            pf.first_line_indent = Emu(327660)
            pf.line_spacing = 1.5
            run = p.add_run(sec.strip())
            run.font.size = Pt(12)
            run.font.name = '仿宋'
    
    doc.save(output_path)
    return {
        'target_dcs': target_dcs,
        'course_dc': course_dc,
        'class_dcs': {k: v['course'] for k, v in class_dcs.items()},
        'output': output_path
    }


if __name__ == '__main__':
    if len(sys.argv) < 5:
        print("Usage: python build_docx_v3.py config.json template.docx xlsx_path output.docx")
        sys.exit(1)
    result = build_docx_v3(sys.argv[1], sys.argv[2], sys.argv[3], sys.argv[4])
    print(json.dumps(result, ensure_ascii=False, indent=2))
