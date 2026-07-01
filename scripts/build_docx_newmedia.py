# -*- coding: utf-8 -*-
"""
build_docx_newmedia.py — 基于新媒体营销模板生成附件7
> 操作时间: 2026-07-01 10:25
> 操作agent: default

用法：python build_docx_newmedia.py config.json output.docx
（模板路径硬编码为新媒体营销模板）
"""
import json, sys, os, math
from docx import Document
from docx.shared import Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

TEMPLATE_PATH = r'C:\Users\Administrator\Desktop\《新媒体营销》试卷分析&教学目标达成情况-2025-2026-2_最终.docx'

def set_cell_text(cell, text, font_name='Times New Roman', font_size=Pt(10.5)):
    """Set cell text WITHOUT clearing — preserves merged cell structure."""
    paragraphs = cell.paragraphs
    if not paragraphs:
        return
    p = paragraphs[0]
    if p.runs:
        p.runs[0].text = str(text)
        for run in p.runs[1:]:
            run.text = ''
    else:
        run = p.add_run(str(text))
        run.font.size = font_size
        run.font.name = font_name

def replace_in_cell(cell, old_text, new_text):
    """Replace old_text with new_text — safe for merged cells (no paragraph clearing)."""
    for p in cell.paragraphs:
        for run in p.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                return True
    return False

def replace_in_cell_fuzzy(cell, old_fragment, new_fragment):
    """Fuzzy replace: searches each run for old_fragment, handles split runs."""
    for p in cell.paragraphs:
        for run in p.runs:
            if old_fragment in run.text:
                run.text = run.text.replace(old_fragment, new_fragment)
                return True
    return False

def replace_teacher_in_row(row, old_teacher, new_teacher):
    """Replace teacher name even when split across runs."""
    for cell in row.cells:
        if replace_in_cell(cell, '任课教师：' + old_teacher, '任课教师：' + new_teacher):
            continue
        replace_in_cell_fuzzy(cell, old_teacher, new_teacher)

def get_distribution(scores):
    n = len(scores)
    if n == 0:
        return {'total':0,'max':0,'min':0,'avg':0,'pass_rate':0,'counts':[0]*7,'pcts':[0]*7}
    max_s, min_s = max(scores), min(scores)
    avg_s = sum(scores) / n
    pass_rate = len([s for s in scores if s >= 60]) / n * 100
    counts = [0]*7
    for s in scores:
        if   s >= 90: counts[0] += 1
        elif s >= 80: counts[1] += 1
        elif s >= 70: counts[2] += 1
        elif s >= 60: counts[3] += 1
        elif s >= 50: counts[4] += 1
        elif s >= 40: counts[5] += 1
        else:         counts[6] += 1
    pcts = [c/n*100 for c in counts]
    return {'total':n,'max':max_s,'min':min_s,'avg':avg_s,'pass_rate':pass_rate,
            'counts':counts,'pcts':pcts}

def replace_all_in_row(row, old_text, new_text):
    """Replace text in ALL cells of a row (handles separate merge groups)."""
    for cell in row.cells:
        replace_in_cell(cell, old_text, new_text)

def replace_all_in_cells(cells, old_text, new_text):
    for cell in cells:
        replace_in_cell(cell, old_text, new_text)

def fill_stat_line(cell, max_v, min_v, avg_v, pass_rate):
    text = f"最高分：{int(max_v)} 最低分：{int(min_v)} 平均分：{avg_v:.1f} 及格率：{pass_rate:.1f}%"
    for p in cell.paragraphs:
        if p.runs:
            p.runs[0].text = text
            for run in p.runs[1:]:
                run.text = ''
            return
    set_cell_text(cell, text)
    text = f"最高分：{int(max_v)} 最低分：{int(min_v)} 平均分：{avg_v:.1f} 及格率：{pass_rate:.1f}%"
    # Replace inline to preserve merged cell structure
    for p in cell.paragraphs:
        if p.runs:
            p.runs[0].text = text
            for run in p.runs[1:]:
                run.text = ''
            return
    # Fallback
    set_cell_text(cell, text)

def remove_vmerge_in_row(row):
    for cell in row.cells:
        tc_pr = cell._tc.find(f'{{{NS_W}}}tcPr')
        if tc_pr is not None:
            for vm in tc_pr.findall(f'{{{NS_W}}}vMerge'):
                tc_pr.remove(vm)

def build_docx(config_path, output_path):
    with open(config_path, 'r', encoding='utf-8') as f:
        cfg = json.load(f)
    
    tg = cfg['targets']
    nt = len(tg)
    students = cfg['students']
    n_all = len(students)
    
    # ── 期末成绩分布 ──
    qm_scores = [st['luru_qm'] for st in students]
    d_all = get_distribution(qm_scores)
    
    class_students = {}
    for st in students:
        cls = st.get('class', cfg['classes'][0])
        class_students.setdefault(cls, []).append(st)
    class_names = list(class_students.keys())
    
    class_dists = {}
    for cls, sts in class_students.items():
        class_dists[cls] = get_distribution([st['luru_qm'] for st in sts])
    
    # ── 达成度计算 ──
    all_item_avgs = {}
    for ti, t in enumerate(tg):
        for ii, it in enumerate(t['items']):
            vals = [st['raw_scores'].get(it['name'],0)/it['raw_full']*it['zhe_full']
                    if it['raw_full'] else 0 for st in students]
            all_item_avgs[(ti, ii)] = round(sum(vals)/len(vals), 2) if vals else 0
    
    target_dcs = {}
    for ti, t in enumerate(tg):
        sum_avg = sum(all_item_avgs[(ti,ii)] for ii in range(len(t['items'])))
        sum_full = sum(it['zhe_full'] for it in t['items'])
        target_dcs[t['name']] = round(sum_avg/sum_full, 2) if sum_full else 0
    
    course_dc = round(sum(target_dcs.values())/len(target_dcs), 2)
    
    class_dcs = {}
    for cls, sts in class_students.items():
        cls_avgs = {}
        for ti, t in enumerate(tg):
            for ii, it in enumerate(t['items']):
                vals = [st['raw_scores'].get(it['name'],0)/it['raw_full']*it['zhe_full']
                        if it['raw_full'] else 0 for st in sts]
                cls_avgs[(ti,ii)] = round(sum(vals)/len(vals), 2) if vals else 0
        cls_tdcs = []
        for ti, t in enumerate(tg):
            sa = sum(cls_avgs[(ti,ii)] for ii in range(len(t['items'])))
            sf = sum(it['zhe_full'] for it in t['items'])
            cls_tdcs.append(round(sa/sf, 2) if sf else 0)
        class_dcs[cls] = {'targets': cls_tdcs, 'course': round(sum(cls_tdcs)/len(cls_tdcs), 2)}
    
    # ═══════ 加载模板 ═══════
    doc = Document(TEMPLATE_PATH)
    t0, t1 = doc.tables[0], doc.tables[1]
    
    # ═══════ Table0 填充（使用 replace_all_in_row 避免合并问题）═══════
    # R0: 课程名称 + 课程代码 + 课程性质
    replace_all_in_row(t0.rows[0], '新媒体营销', cfg['course_name'])
    replace_all_in_row(t0.rows[0], '3300000008', cfg['course_code'])
    replace_all_in_row(t0.rows[0], '通识教育必修', '必修课（通识教育课）')
    
    # R1: 开课部门 + 考核方式 + 学分
    replace_all_in_row(t0.rows[1], '2', str(cfg.get('credits','2')))
    
    # R2: 教学班 + 应考人数 + 实考人数
    all_cls_str = '、'.join(c + '班' for c in cfg.get('classes', []))
    for cell in t0.rows[2].cells:
        replace_in_cell(cell, '计科243、244班', all_cls_str)
        replace_in_cell(cell, '应考人数：79', '应考人数：' + str(n_all))
        replace_in_cell(cell, '实考人数：79', '实考人数：' + str(n_all))
    
    # R4-R5: 分数段人数和百分比（非合并单元格，set_cell_text 安全）
    score_cols = [2, 3, 5, 6, 7, 9, 10]
    dup_map = {3: 4, 7: 8}
    for ci, col in enumerate(score_cols):
        set_cell_text(t0.rows[4].cells[col], str(d_all['counts'][ci]))
        set_cell_text(t0.rows[5].cells[col], f"{d_all['pcts'][ci]:.1f}%")
        if col in dup_map:
            set_cell_text(t0.rows[4].cells[dup_map[col]], str(d_all['counts'][ci]))
            set_cell_text(t0.rows[5].cells[dup_map[col]], f"{d_all['pcts'][ci]:.1f}%")
    
    # R6: 统计行 — 写入 C1（gridSpan=10，不是C0）
    fill_stat_line(t0.rows[6].cells[1], d_all['max'], d_all['min'], d_all['avg'], d_all['pass_rate'])
    
    # 分班: R7-R12 (class1), R13-R17 (class2)
    # 注: 模板 R8(计科243)=41人, R13(计科244)=38人 — 先dump确认！
    class_template = {
        0: {'base_r': 8, 'old_cls': '计科243班', 'old_teacher': '刘娓娓', 'old_count': '41'},
        1: {'base_r': 13, 'old_cls': '计科244班', 'old_teacher': '刘娓娓', 'old_count': '38'}
    }
    
    for ci in range(min(len(class_names), 2)):
        cname = class_names[ci]
        dist = class_dists[cname]
        ct = class_template[ci]
        base_r = ct['base_r']
        new_teacher = cfg.get('teacher', '邹芳')
        
        # R(base): 替换教学班/教师/人数
        row_title = t0.rows[base_r]
        for cell in row_title.cells:
            replace_in_cell(cell, ct['old_cls'], cname + '班')
            replace_in_cell(cell, '应考人数：' + ct['old_count'], '应考人数：' + str(dist['total']))
            replace_in_cell(cell, '实考人数：' + ct['old_count'], '实考人数：' + str(dist['total']))
        replace_teacher_in_row(row_title, ct['old_teacher'], new_teacher)
        
        # R(base+2): 人数, R(base+3): 百分比
        for cj, col in enumerate(score_cols):
            set_cell_text(t0.rows[base_r+2].cells[col], str(dist['counts'][cj]))
            set_cell_text(t0.rows[base_r+3].cells[col], f"{dist['pcts'][cj]:.1f}%")
            if col in dup_map:
                set_cell_text(t0.rows[base_r+2].cells[dup_map[col]], str(dist['counts'][cj]))
                set_cell_text(t0.rows[base_r+3].cells[dup_map[col]], f"{dist['pcts'][cj]:.1f}%")
        
        # R(base+4): 统计行
        fill_stat_line(t0.rows[base_r+4].cells[1], dist['max'], dist['min'], dist['avg'], dist['pass_rate'])
    
    # 单班情况：将第二个班级区域也填成同一个班的数据
    if len(class_names) == 1:
        cname = class_names[0]
        dist = class_dists[cname]
        ct = class_template[1]
        new_teacher = cfg.get('teacher', '邹芳')
        
        for cell in t0.rows[ct['base_r']].cells:
            replace_in_cell(cell, ct['old_cls'], cname + '班')
            replace_in_cell(cell, '应考人数：' + ct['old_count'], '应考人数：' + str(dist['total']))
            replace_in_cell(cell, '实考人数：' + ct['old_count'], '实考人数：' + str(dist['total']))
        replace_teacher_in_row(t0.rows[ct['base_r']], ct['old_teacher'], new_teacher)
        
        for cj, col in enumerate(score_cols):
            set_cell_text(t0.rows[ct['base_r']+2].cells[col], str(dist['counts'][cj]))
            set_cell_text(t0.rows[ct['base_r']+3].cells[col], f"{dist['pcts'][cj]:.1f}%")
            if col in dup_map:
                set_cell_text(t0.rows[ct['base_r']+2].cells[dup_map[col]], str(dist['counts'][cj]))
                set_cell_text(t0.rows[ct['base_r']+3].cells[dup_map[col]], f"{dist['pcts'][cj]:.1f}%")
        
        fill_stat_line(t0.rows[ct['base_r']+4].cells[1], dist['max'], dist['min'], dist['avg'], dist['pass_rate'])
    
    # ── Table1 原位覆写（不删行不增行，保留模板合并结构）──
    # 原模板: R0(header), R1-R24(data 4×6), R25(分析), R26(签名), R27(空)
    # 大学计算机基础: 4×5=20 → 前5行填数据，第6行清空
    
    for ti, t in enumerate(tg):
        for ii in range(6):  # 模板每目标6行（5项+项目作业）
            row_idx = ti * 6 + ii + 1  # R1, R2, ..., R24
            row = t1.rows[row_idx]
            
            if ii >= len(t['items']):
                # 第6行（原项目作业）→ 清空内容（C9除外，模板已按目标合并）
                for ci in [5, 6, 7, 8]:  # 考核环节+分值+平均成绩（C9由合并自动填充）
                    set_cell_text(row.cells[ci], '—')
                continue
            
            it = t['items'][ii]
            
            # C0: 课程教学目标达成计算 (模板已合并，只在首行写)
            if ii == 0 and ti == 0:
                set_cell_text(row.cells[0], '课程教学目标达成计算')
            
            # C1: 毕业要求（模板已合并，只在首行写）
            if ii == 0:
                set_cell_text(row.cells[1], t.get('grad_req', '毕业要求' + t.get('indicator','').split('-')[0]))
                set_cell_text(row.cells[2], t.get('indicator', ''))
                set_cell_text(row.cells[3], t['name'])
            
            # C5: 考核环节
            set_cell_text(row.cells[5], it['name'])
            # C6: 分值(A) raw_full
            set_cell_text(row.cells[6], str(it['raw_full']))
            # C7: 分值(A) zhe_full  
            set_cell_text(row.cells[7], str(it['zhe_full']))
            # C8: 平均成绩(B)
            avg_val = all_item_avgs.get((ti, ii), 0)
            set_cell_text(row.cells[8], str(avg_val))
            # C9: 达成度（模板按目标合并，每目标一个值）
            if ii == 0:
                tdc = target_dcs[f'教学目标{ti+1}']
                set_cell_text(row.cells[9], f'{tdc:.2f}')
    
    # ═══════ 分析文本 ═══════
    classes_desc = '、'.join(c + '班' for c in cfg.get('classes', []))
    below_60 = 100 - d_all['pass_rate']
    below_50 = d_all['pcts'][5] + d_all['pcts'][6]
    
    class_comp = ''
    if len(class_names) >= 2:
        c1, c2 = class_names[0], class_names[1]
        dc1, dc2 = class_dcs[c1]['course'], class_dcs[c2]['course']
        diff = abs(dc1 - dc2)
        class_comp = f"\n\n分班达成度对比：{c1}班={dc1:.2f}，{c2}班={dc2:.2f}，差异为{diff:.2f}。"
        if diff > 0.05:
            class_comp += f" {c2}班达成度明显低于{c1}班，需重点关注。"
    
    item_rates = {}
    for ti, t in enumerate(tg):
        for ii, it in enumerate(t['items']):
            avg = all_item_avgs.get((ti,ii),0)
            rate = avg/it['zhe_full'] if it['zhe_full'] else 0
            item_rates[(ti,ii)] = {'tname':t['name'],'iname':it['name'],'avg':avg,'full':it['zhe_full'],'rate':rate}
    weakest = sorted(item_rates.values(), key=lambda x: x['rate'])[:3]
    target_ranking = sorted(target_dcs.items(), key=lambda x: x[1])
    
    low_students = [st for st in students if st.get('luru_qm',0) < 60]
    low_names = '、'.join(st['name'] for st in low_students[:5])
    if len(low_students) > 5: low_names += f'等{len(low_students)}人'
    
    targets_dc_text = '，'.join([f'{name}={val:.2f}' for name, val in target_dcs.items()])
    
    analysis_text = f"""一、终结性考核典型性错误及试题难度分析

（一）终结性考核概况
本课程终结性考核采用期末考试形式（闭卷笔试，满分100分），面向{classes_desc}共{n_all}名学生。考试最高分{int(d_all['max'])}分，最低分{int(d_all['min'])}分，平均分{d_all['avg']:.1f}分，及格率{d_all['pass_rate']:.1f}%。从整体成绩分布来看，考试成绩呈现"{"两极分化较明显" if int(d_all['max'])-int(d_all['min']) >= 40 else "分布相对集中"}"的特征。

（二）成绩分段分布
90-100分{d_all['counts'][0]}人（{d_all['pcts'][0]:.1f}%），80-89分{d_all['counts'][1]}人（{d_all['pcts'][1]:.1f}%），70-79分{d_all['counts'][2]}人（{d_all['pcts'][2]:.1f}%），60-69分{d_all['counts'][3]}人（{d_all['pcts'][3]:.1f}%），50-59分{d_all['counts'][4]}人（{d_all['pcts'][4]:.1f}%），40-49分{d_all['counts'][5]}人（{d_all['pcts'][5]:.1f}%），40分以下{d_all['counts'][6]}人（{d_all['pcts'][6]:.1f}%）。{"及格率超过85%，整体教学效果良好。" if d_all['pass_rate'] >= 85 else ("及格率处于中等水平，存在一定提升空间。" if d_all['pass_rate'] >= 70 else "不及格比例偏高，需引起高度重视。")}

（三）典型性错误分析
从考试结果看，{below_60:.0f}%的学生成绩低于60分{"，低于50分的有" + str(int(below_50)) + "%" if below_50 > 0 else ""}。不及格学生{"包括" + low_names if low_students else ""}，反映出部分学生在基础概念理解、计算准确性及综合应用等方面存在不足。建议在阅卷中对典型性错误进行分类记录，以便后续针对性纠错。

二、课程教学目标达成分析

本课程共{nt}个教学目标，对应课程教学质量标准中的毕业要求指标点。各目标达成度如下：{targets_dc_text}。课程整体达成度为{course_dc:.2f}，处于{"较高" if course_dc >= 0.75 else ("中等" if course_dc >= 0.65 else "偏低")}水平，{"教学目标得以较好实现" if course_dc >= 0.75 else ("教学目标基本达成，部分需加强" if course_dc >= 0.65 else "部分目标达成度偏低，亟待改进")}。

达成度最高的是{target_ranking[-1][0]}（{target_ranking[-1][1]:.2f}），最低的是{target_ranking[0][0]}（{target_ranking[0][1]:.2f}）。得分率最低的三个考核环节为：{"；".join(f'"{w["tname"]}-{w["iname"]}"达成率{w["rate"]:.1%}' for w in weakest)}。{class_comp}

三、学生学习和教师教学过程中存在的问题

（一）学生学习方面
1. 基础知识掌握不牢固。部分学生在{weakest[0]['iname']}中得分偏低（达成率仅{weakest[0]['rate']:.1%}），反映基础概念和基本技能掌握不扎实。
2. 综合应用能力有待提升。{"期末考试成绩分布跨度大（" + str(int(d_all["max"])) + "分-" + str(int(d_all["min"])) + "分），" if int(d_all['max'])-int(d_all['min']) >= 30 else ""}学生综合运用知识解决问题的能力参差不齐。
3. 自主学习意识不足。{"不及格学生" + str(len(low_students)) + "人（占" + str(round(len(low_students)/n_all*100,1)) + "%），" if low_students else ""}部分学生未在课程过程中及时巩固知识。
4. {"平时成绩与期末成绩的关联性需关注，部分学生平时考核表现较好但期末不佳，平时考核可能未真实反映学习效果。" if d_all['pass_rate'] < 90 else ""}

（二）教师教学方面
1. 薄弱环节专项训练不足，达成度低的环节集中在{"、".join(w['iname'] for w in weakest[:2])}等方面。
2. 过程性帮扶机制有待加强，对学情预警响应不够及时。
3. 教学资源布局可进一步优化，特别是薄弱环节的补充材料。
4. 不同班级教学效果的均衡性需关注，应统一标准并加强薄弱班督导。

四、对今后教学的改进意见

1. 强化薄弱环节专项训练。针对{weakest[0]['tname']}的"{weakest[0]['iname']}"（达成率{weakest[0]['rate']:.1%}），设计专项练习和阶段性测验。
2. 建立学业预警与帮扶机制。定期分析学生成绩数据，对不及格风险学生进行重点关注与课后辅导。
3. 优化教学资源配置。录制薄弱知识点精讲视频，建设在线自测题库。
4. 加强过程性考核诊断功能。学期中增加模拟测验，阶段测试结果进行知识点归因分析。
5. 推进班级教学均衡。{"统一教学进度和考核标准，对薄弱班级增加互动与随堂测验。" if len(class_names) >= 2 else "持续关注全班学习动态，对后进生实行一对一帮扶。"}"""
    
    # 写入分析区（原模板 R25，保持不变）
    analysis_row_idx = 25
    targets_dc_text = '，'.join([f'{name}={val:.2f}' for name, val in target_dcs.items()])
    
    dc_header = f"课程整体达成度为{course_dc:.2f}（{targets_dc_text}），处于{'较高' if course_dc >= 0.75 else ('中等' if course_dc >= 0.65 else '偏低')}水平。\n\n"
    
    analysis_text = dc_header + analysis_text
    if analysis_row_idx < len(t1.rows):
        cell = t1.rows[analysis_row_idx].cells[1]
        for p in cell.paragraphs: p.clear()
        sections = analysis_text.split('\n\n')
        for i, sec in enumerate(sections):
            if i == 0:
                p = cell.paragraphs[0]
            else:
                p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf = p.paragraph_format
            pf.first_line_indent = Emu(327660)
            pf.line_spacing = 1.5
            run = p.add_run(sec.strip())
            run.font.size = Pt(12)
            run.font.name = '仿宋'
    
    doc.save(output_path)
    return {'target_dcs': target_dcs, 'course_dc': course_dc,
            'class_dcs': {k: v['course'] for k, v in class_dcs.items()}, 'output': output_path}

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python build_docx_newmedia.py config.json output.docx")
        sys.exit(1)
    result = build_docx(sys.argv[1], sys.argv[2])
    print(json.dumps(result, ensure_ascii=False, indent=2))
