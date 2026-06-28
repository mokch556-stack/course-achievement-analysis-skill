#!/usr/bin/env python3
"""
scan_sources.py — 数据源全量扫描脚本
功能：扫描指定目录下所有相关文件，识别每个文件的结构
输出：JSON格式报告，列出找到的xlsx/docx及各自的sheet/表格结构

操作时间: 2026-06-28
操作agent: default
"""

import sys
import json
import argparse
from pathlib import Path

try:
    import openpyxl
except ImportError:
    print("❌ 需要 openpyxl 库")
    sys.exit(1)

try:
    from docx import Document
except ImportError:
    print("❌ 需要 python-docx 库")
    sys.exit(1)


def scan_xlsx(path):
    """扫描一个xlsx文件，返回其所有sheet的结构信息"""
    try:
        wb = openpyxl.load_workbook(path, data_only=True)
    except Exception as e:
        return {"path": str(path), "error": str(e), "sheets": []}

    result = {"path": str(path), "error": None, "sheets": []}
    
    for name in wb.sheetnames:
        ws = wb[name]
        sheet_info = {
            "name": name,
            "rows": ws.max_row,
            "cols": ws.max_column,
            "header_rows": []
        }
        # 读取前10行作为结构参考
        for r in range(1, min(ws.max_row + 1, 11)):
            row_vals = []
            for c in range(1, min(ws.max_column + 1, 15)):  # 最多读15列
                v = ws.cell(r, c).value
                if isinstance(v, float):
                    v = round(v, 2)
                row_vals.append(str(v) if v is not None else "")
            sheet_info["header_rows"].append(row_vals)
        
        # 识别关键列
        # 读取第8行（数据开始行）的前几列
        if ws.max_row >= 8:
            sample_row = []
            for c in range(1, min(ws.max_column + 1, 15)):
                v = ws.cell(8, c).value
                if isinstance(v, float):
                    v = round(v, 2)
                sample_row.append(str(v) if v is not None else "")
            sheet_info["sample_data_row"] = sample_row
        
        result["sheets"].append(sheet_info)
    
    return result


def scan_docx(path):
    """扫描一个docx文件，返回其表格结构信息"""
    try:
        doc = Document(path)
    except Exception as e:
        return {"path": str(path), "error": str(e), "tables": [], "paragraphs": 0}

    result = {
        "path": str(path),
        "error": None,
        "paragraphs": len(doc.paragraphs),
        "tables": []
    }
    
    for i, t in enumerate(doc.tables):
        table_info = {
            "index": i,
            "rows": len(t.rows),
            "cols": len(t.columns),
            "header_rows": []
        }
        for j, row in enumerate(t.rows[:8]):  # 前8行
            cells = [c.text[:30] for c in row.cells]
            table_info["header_rows"].append(cells)
        result["tables"].append(table_info)
    
    return result


def find_relevant_files(directory):
    """在目录中查找所有与达成度分析相关的文件"""
    d = Path(directory)
    if not d.exists():
        return {"error": f"目录不存在: {directory}", "files": []}

    result = {"directory": str(d.absolute()), "files": []}
    
    # 定义关键词权重
    score_keywords = [
        ("教学大纲", 10), ("开课说明", 10),
        ("成绩", 8), ("登记表", 7), ("总评", 6),
        ("平时", 6), ("期末", 6),
        ("达成度", 7), ("分析报告", 7),
        ("模板", 3), ("附件", 4),
        ("试卷", 5), ("考试", 5),
    ]
    
    for f in sorted(d.rglob("*")):
        if not f.is_file():
            continue
        ext = f.suffix.lower()
        if ext not in ('.xlsx', '.xls', '.xlsm', '.docx', '.pdf'):
            continue
        
        # 跳过废弃/中间文件
        name = f.name.lower()
        if any(x in name for x in ('废弃', '中间', '备份', 'backup', '~$')):
            continue
        
        # 计算相关性分数
        score = 0
        for keyword, weight in score_keywords:
            if keyword in name:
                score += weight
        
        # 只保留相关文件（score >= 3 或 xlsx标准格式）
        if score < 3:
            continue
        
        file_info = {
            "name": f.name,
            "path": str(f.absolute()),
            "size": f.stat().st_size,
            "ext": ext,
            "relevance_score": score
        }
        
        # 扫描内容
        if ext in ('.xlsx', '.xls', '.xlsm'):
            try:
                scan = scan_xlsx(f)
                file_info["sheets"] = scan["sheets"]
                file_info["total_sheets"] = len(scan.get("sheets", []))
            except:
                file_info["error"] = "扫描失败"
        elif ext == '.docx':
            try:
                scan = scan_docx(f)
                file_info["tables"] = scan["tables"]
                file_info["total_paragraphs"] = scan["paragraphs"]
            except:
                file_info["error"] = "扫描失败"
        
        result["files"].append(file_info)
    
    # 按相关性分数排序
    result["files"].sort(key=lambda x: x.get("relevance_score", 0), reverse=True)
    
    return result


def main():
    parser = argparse.ArgumentParser(description='数据源全量扫描脚本')
    parser.add_argument('directory', help='要扫描的目录路径')
    parser.add_argument('--output', '-o', help='输出JSON文件路径（默认打印到控制台）')
    args = parser.parse_args()

    result = find_relevant_files(args.directory)
    
    if args.output:
        with open(args.output, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        print(f"✅ 扫描报告已保存: {args.output}")
    else:
        print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == '__main__':
    main()
