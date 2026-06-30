# -*- coding: utf-8 -*-
"""
附件7 v6.2 — 修复：课程性质/应考实考/C5-C8/算术平均/删图2
"""
import os, openpyxl, io, matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'SimSun']
plt.rcParams['axes.unicode_minus'] = False

TMPL = r"C:\Users\Administrator\Desktop\齐工程教〔2026〕64号关于2025-2026学年度第二学期期末考试工作安排的通知\齐工程教〔2026〕64号关于2025-2026学年度第二学期期末考试工作安排的通知\附件7.齐齐哈尔工程学院课程教学目标达成情况分析报告.docx"
OUT_PATH = r"E:\qwenpaw\workspaces\default\附件7_农产品供应链_教学目标达成情况分析报告_v6_3.docx"
CJB1 = r"C:\Users\Administrator\Desktop\农产品供应链\课程档案\供应链1班成绩登记表.xlsx"
CJB2 = r"C:\Users\Administrator\Desktop\农产品供应链\课程档案\供应链2班成绩登记表.xlsx"
XLSX = r"E:\qwenpaw\workspaces\default\skills\course_achievement_analysis\scripts\达成度计算表_农产品供应链_电商231班+电商232班_v4.xlsx"

COURSE_CODE = "3207000182"
COURSE_NATURE = "选修课"
EXAM_MODE = "考查"
TEACHER = "赵淑雯"
HOURS = 16
CREDITS = 1

def fill_label_cell(cell, append_val):
    """追加内容到单元格末尾，保留模板原有标签文字"""
    text = cell.text.strip()
    if str(append_val) in text:
        return  # 已存在，跳过
    para = cell.paragraphs[0]
    if not para.runs:
        para.add_run(str(append_val))
    elif para.runs[-1].text.strip() == '':
        para.runs[-1].text = str(append_val)
    elif str(append_val) not in para.runs[-1].text:
        para.runs[-1].text += str(append_val)
    # 设置字体
    for run in para.runs:
        run.font.size = Pt(10.5)
        run.font.name = 'Times New Roman'

def set_cell_text_clean(cell, text, font_name='Times New Roman', font_size=10.5):
    for p in cell.paragraphs:
        p.clear()
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(font_size)
    run.font.name = font_name

def get_students():
    students = []
    for path, cls in [(CJB1, '电商231班'), (CJB2, '电商232班')]:
        wb = openpyxl.load_workbook(path, data_only=True)
        ws = wb.active
        for r in range(7, ws.max_row + 1):
            name = ws.cell(r, 3).value
            qm = ws.cell(r, 6).value
            if name and qm is not None:
                students.append({'name': str(name).strip(), 'cls': cls, 'qm': float(qm)})
        wb.close()
    return students

def get_distribution(scores):
    n = len(scores)
    max_s, min_s = max(scores), min(scores)
    avg_s = sum(scores) / n
    pass_rate = len([s for s in scores if s >= 60]) / n * 100
    segments = [(90, 100), (80, 89), (70, 79), (60, 69), (50, 59), (30, 49), (0, 29)]
    counts = [len([s for s in scores if lo <= s <= hi]) for lo, hi in segments]
    return {'total': n, 'max': max_s, 'min': min_s, 'avg': round(avg_s, 1),
            'pass_rate': round(pass_rate, 1), 'counts': counts,
            'pcts': [c/n*100 for c in counts]}

def fill_stat_row(row, max_v, min_v, avg_v, pass_rate):
    """填充R6型统计行（合并单元格，含「最高分：」「最低分：」等标签）"""
    cell = row.cells[1]
    para = cell.paragraphs[0]
    for r in para.runs:
        r._element.getparent().remove(r._element)
    text = f"最高分：{int(max_v)}        最低分：{int(min_v)}        平均分：{avg_v}        及格率：{pass_rate}%"
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10.5)

def append_analysis_para(cell, text):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.first_line_indent = Emu(327660)
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '仿宋'
    return p

print("[1] 读取成绩登记表")
students = get_students()
scores_total = [s['qm'] for s in students]
scores_231 = [s['qm'] for s in students if s['cls'] == '电商231班']
scores_232 = [s['qm'] for s in students if s['cls'] == '电商232班']
n_all = len(students)
n_231 = len(scores_231)
n_232 = len(scores_232)
print(f"  学生总数={n_all}, 电商231={n_231}, 电商232={n_232}")

print("[2] 从xlsx读取平均分")
wb_xlsx = openpyxl.load_workbook(XLSX)
ws_xlsx = wb_xlsx.active
xlsx_totals = [0]*10
xlsx_cnt = 0
for r in range(10, 74):
    vals = [ws_xlsx.cell(r, c).value for c in range(4, 14)]
    if vals[0] is not None and isinstance(vals[0], (int, float)):
        for i in range(10):
            xlsx_totals[i] += (vals[i] if isinstance(vals[i], (int, float)) else 0)
        xlsx_cnt += 1
avg_g1_ps = round(sum(xlsx_totals[0:4]) / xlsx_cnt, 2) if xlsx_cnt else 0
avg_g1_qm = round(xlsx_totals[4] / xlsx_cnt, 2) if xlsx_cnt else 0
avg_g2_ps = round(sum(xlsx_totals[5:9]) / xlsx_cnt, 2) if xlsx_cnt else 0
avg_g2_qm = round(xlsx_totals[9] / xlsx_cnt, 2) if xlsx_cnt else 0
wb_xlsx.close()
print(f"  目标1：平时平均={avg_g1_ps}, 期末平均={avg_g1_qm}")
print(f"  目标2：平时平均={avg_g2_ps}, 期末平均={avg_g2_qm}")

print("[3] 计算成绩分布")
d_all = get_distribution(scores_total)
d_231 = get_distribution(scores_231)
d_232 = get_distribution(scores_232)
print(f"  全级: {d_all['total']}人, 最高={d_all['max']}, 最低={d_all['min']}, "
      f"平均={d_all['avg']}, 及格率={d_all['pass_rate']}%")

print("[4] 读取达成度")
wb = openpyxl.load_workbook(XLSX, data_only=True)
ws = wb.active
r76_text = str(ws.cell(76, 1).value or '')
import re
m1 = re.search(r'目标1=([\d.]+)', r76_text)
m2 = re.search(r'目标2=([\d.]+)', r76_text)
g1 = round(float(m1.group(1)), 2) if m1 else 0.0
g2 = round(float(m2.group(1)), 2) if m2 else 0.0
cv = round((g1 + g2) / 2, 2)
wb.close()
print(f"  目标1={g1}, 目标2={g2}, 课程={cv} (算术平均)")

print("[5] 加载模板")
doc = Document(TMPL)
t0 = doc.tables[0]
t1 = doc.tables[1]

# ======== Table0 ========
print("--- 填充 Table0 ---")

# R0: 课程基本信息
fill_label_cell(t0.rows[0].cells[0], "农产品供应链")
fill_label_cell(t0.rows[0].cells[5], COURSE_CODE)
fill_label_cell(t0.rows[0].cells[8], COURSE_NATURE)

# R1: 开课部门/考核方式/学分
fill_label_cell(t0.rows[1].cells[0], "信息学院")
fill_label_cell(t0.rows[1].cells[5], EXAM_MODE)
fill_label_cell(t0.rows[1].cells[8], CREDITS)

# R2: 教学班/应考/实考
fill_label_cell(t0.rows[2].cells[1], "电商231、电商232")
fill_label_cell(t0.rows[2].cells[6], str(n_all))
fill_label_cell(t0.rows[2].cells[9], str(n_all))

# R3: 分数段标签行 → 保留模板原样，不动

# R4: 人数行（全级）
col_keys = [2, 3, 5, 6, 7, 9, 10]
for ci, cnt in enumerate(d_all['counts']):
    fill_label_cell(t0.rows[4].cells[col_keys[ci]], str(cnt))

# R5: 百分比行（全级）
for ci, pct in enumerate(d_all['pcts']):
    fill_label_cell(t0.rows[5].cells[col_keys[ci]], f"{pct:.1f}%")

# R6: 统计行（全级 - 合并行）
fill_stat_row(t0.rows[6], d_all['max'], d_all['min'], d_all['avg'], d_all['pass_rate'])

# R8: 行政班1（电商231）
fill_label_cell(t0.rows[8].cells[1], "电商231")
fill_label_cell(t0.rows[8].cells[4], TEACHER)
fill_label_cell(t0.rows[8].cells[6], str(n_231))
fill_label_cell(t0.rows[8].cells[9], str(n_231))

# R9: 分数段标签行（行政班1）→ 不动
# R10: 人数行（行政班1）
for ci, cnt in enumerate(d_231['counts']):
    fill_label_cell(t0.rows[10].cells[col_keys[ci]], str(cnt))
# R11: 百分比行（行政班1）
for ci, pct in enumerate(d_231['pcts']):
    fill_label_cell(t0.rows[11].cells[col_keys[ci]], f"{pct:.1f}%")
# R12: 统计行（行政班1 - 合并行）
fill_stat_row(t0.rows[12], d_231['max'], d_231['min'], d_231['avg'], d_231['pass_rate'])

# R13: 行政班2（电商232）
fill_label_cell(t0.rows[13].cells[1], "电商232")
fill_label_cell(t0.rows[13].cells[4], TEACHER)
fill_label_cell(t0.rows[13].cells[6], str(n_232))
fill_label_cell(t0.rows[13].cells[9], str(n_232))

# R14: 分数段标签行（行政班2）→ 不动
# R15: 人数行（行政班2）
for ci, cnt in enumerate(d_232['counts']):
    fill_label_cell(t0.rows[15].cells[col_keys[ci]], str(cnt))
# R16: 百分比行（行政班2）
for ci, pct in enumerate(d_232['pcts']):
    fill_label_cell(t0.rows[16].cells[col_keys[ci]], f"{pct:.1f}%")
# R17: 统计行（行政班2 - 合并行）
fill_stat_row(t0.rows[17], d_232['max'], d_232['min'], d_232['avg'], d_232['pass_rate'])
print("[OK] Table0 filled")

# ======== Table1 ========
print("--- 处理 Table1 ---")
tbl1 = t1._tbl
tbl1.remove(tbl1[5])
print("[OK] Deleted R5 (目标3)")
set_cell_text_clean(t1.rows[3].cells[1], "毕业要求1")
set_cell_text_clean(t1.rows[3].cells[2], "1-2")
set_cell_text_clean(t1.rows[3].cells[3], "课程教学目标2")
set_cell_text_clean(t1.rows[3].cells[4], "课程教学目标2")
set_cell_text_clean(t1.rows[4].cells[1], "毕业要求1")
set_cell_text_clean(t1.rows[4].cells[2], "1-2")
set_cell_text_clean(t1.rows[4].cells[3], "课程教学目标2")
set_cell_text_clean(t1.rows[4].cells[4], "课程教学目标2")

# C9达成度
set_cell_text_clean(t1.rows[1].cells[9], f"{g1:.2f}")
set_cell_text_clean(t1.rows[2].cells[9], f"{g1:.2f}")
set_cell_text_clean(t1.rows[3].cells[9], f"{g1:.2f}")
set_cell_text_clean(t1.rows[4].cells[9], f"{g2:.2f}")
set_cell_text_clean(t1.rows[5].cells[9], f"{cv:.2f}")

# C5-C8考核环节/分值/平均成绩
set_cell_text_clean(t1.rows[1].cells[5], "过程性考核")
set_cell_text_clean(t1.rows[1].cells[6], "25")
set_cell_text_clean(t1.rows[1].cells[7], "")
set_cell_text_clean(t1.rows[1].cells[8], f"{avg_g1_ps}")
set_cell_text_clean(t1.rows[2].cells[5], "期末项目")
set_cell_text_clean(t1.rows[2].cells[6], "20")
set_cell_text_clean(t1.rows[2].cells[7], "")
set_cell_text_clean(t1.rows[2].cells[8], f"{avg_g1_qm}")
avg_g2_total = round(avg_g2_ps + avg_g2_qm, 2)
set_cell_text_clean(t1.rows[4].cells[5], "过程性考核+期末项目")
set_cell_text_clean(t1.rows[4].cells[6], "55")
set_cell_text_clean(t1.rows[4].cells[7], "")
set_cell_text_clean(t1.rows[4].cells[8], f"{avg_g2_total}")
print("[OK] Table1 filled")

# ======== 分析文本 ========
print("--- 生成分析文本 ---")
ac = t1.rows[6].cells[1]

max_students = [s['name'] for s in students if s['qm'] == d_all['max']]
min_students = [s['name'] for s in students if s['qm'] == d_all['min']]
max_names = '、'.join(max_students)
min_names = '、'.join(min_students)
pass_rate_val = round(d_all['pass_rate'], 1)
fail_rate_val = round(100 - pass_rate_val, 1)

text_p1 = (
    "（一）终结性考核典型性错误\n"
    "《农产品供应链》课程期末考核采取项目作业形式，"
    "以农产品供应链基础理论、流通模式、电子商务运营与金融支持为核心内容，"
    f"全面考查学生对课程核心知识的掌握程度。本次考试共有{n_all}名学生参加，"
    f"最高分{d_all['max']}分（{max_names}），"
    f"最低分{d_all['min']}分（{min_names}），"
    f"平均分{d_all['avg']}分，及格率{pass_rate_val}%，"
    f"不及格率{fail_rate_val}%。"
    "从成绩分布来看，大多数学生能够掌握农产品供应链的基础理论和流通模式，"
    "但在电子商务运营与金融支持的深度应用方面存在薄弱环节。"
)

text_p2 = (
    "（二）过程性考核分析\n"
    "过程性考核采用阶段性测试、课堂笔记、项目方案、课程竞赛相结合的方式，"
    "全面覆盖课程各知识点。"
    "从平时成绩看，学生整体出勤率高，课堂参与度良好，"
    "阶段性测试平均分保持在较好水平，说明学生能够及时巩固课堂所学内容。"
    "项目方案环节中，大多数学生能够运用所学知识设计农产品供应链优化方案，"
    "但在数据分析深度和方案创新性方面仍有提升空间。"
    "课堂笔记环节反映出学生对基础知识点的记录较为完整，"
    "但对拓展性知识的归纳总结能力有待加强。"
    "课程竞赛环节激发了学生的学习积极性，"
    "部分学生表现出较强的团队协作能力和创新思维，"
    "整体过程性考核成绩分布合理，能够客观反映学生的平时学习状态。"
)

text_p3 = (
    "（三）教学目标达成情况分析\n"
    f"课程教学目标1达成度{g1}，教学目标2达成度{g2}，"
    f"总达成度{cv}。"
    "两个教学目标达成度均处于良好水平，说明学生经过本课程学习，"
    "基本掌握了农产品供应链的基础理论和实践技能。"
    "教学目标1（掌握农产品供应链基础理论与流通模式）达成度相对较高，"
    "因为该部分内容基础性强，理论与实践结合紧密，学生理解较为深入。"
    "教学目标2（掌握农产品电子商务运营与金融支持）达成度略低，"
    "主要原因是供应链金融部分涉及较多财务管理知识，"
    "学生对跨学科知识的综合应用能力有待提高。"
    "从各考核环节来看，过程性考核中学生的项目方案设计能力和团队协作能力"
    "表现良好，期末项目作业中部分学生对供应链信息技术的应用不够熟练，"
    "建议在今后的教学中增加信息技术与供应链管理融合的实践案例，"
    "加强学生对大数据、物联网等新技术在农产品供应链中应用的理解。"
)

text_p4 = (
    "（四）持续改进措施\n"
    "1. 优化教学内容：根据达成度分析结果，适当增加目标2中供应链金融"
    "和电子商务运营的案例教学比例，引入更多行业真实案例，"
    "强化学生对复杂问题的分析与解决能力，同时精简基础理论部分的课时占比。\n"
    "2. 完善考核方式：进一步优化过程性考核的阶段性测试内容，"
    "增加综合性开放性题目，减少记忆性题目比例，"
    "注重考查学生的分析能力和创新思维，适当提高项目方案在平时成绩中的权重。\n"
    "3. 加强实践教学：增加供应链模拟实训环节，"
    "利用虚拟仿真平台让学生体验农产品从生产到销售的完整供应链流程，"
    "提升学生解决实际问题的能力，同时鼓励学生参与校企合作实践项目。\n"
    "4. 个性化辅导：对达成度偏低的学生进行针对性辅导，"
    "建立学习帮扶机制，安排学有余力的学生带动学习困难的学生，"
    "形成互帮互助的良好学习氛围，确保全体学生的共同进步。"
)

para_map = {1: text_p1, 3: text_p2, 6: text_p3}
for pi in para_map:
    p = ac.paragraphs[pi]
    for run in p.runs:
        run._element.getparent().remove(run._element)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.first_line_indent = Emu(327660)
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    run = p.add_run(para_map[pi])
    run.font.size = Pt(12)
    run.font.name = '仿宋'

append_analysis_para(ac, text_p4)
print("[OK] Analysis text filled")

# ======== 图1: 成绩分布直方图 ========
print("--- 生成图1 ---")
fig, ax = plt.subplots(figsize=(8, 4))
seg_labels = ['90~100', '80~89', '70~79', '60~69', '50~59', '30~49', '0~29']
bars = ax.bar(seg_labels, d_all['counts'], color='steelblue', width=0.6)
for bar, cnt in zip(bars, d_all['counts']):
    ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.3,
            str(cnt), ha='center', va='bottom', fontsize=10)
ax.set_xlabel('分数段', fontsize=11)
ax.set_ylabel('人数', fontsize=11)
ax.set_title('终结性考核成绩分布', fontsize=13)
ax.set_ylim(0, max(d_all['counts']) * 1.2)
chart1_path = r"E:\qwenpaw\workspaces\default\_chart1_dist.png"
plt.tight_layout()
plt.savefig(chart1_path, dpi=200)
plt.close()
print("[OK] 图1 generated")

# 插入图1
body = doc.element.body
for i, para in enumerate(doc.paragraphs):
    if para.text.strip().startswith('（一）终结性考核'):
        p_elem = para._element
        img_p = doc.add_paragraph()
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_p.add_run()
        run.add_picture(chart1_path, width=Inches(5.5))
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_p.add_run("图1  终结性考核成绩分布图")
        cap_run.font.size = Pt(10)
        body.remove(cap_p._element)
        body.remove(img_p._element)
        p_elem.addprevious(cap_p._element)
        cap_p._element.addprevious(img_p._element)
        break
print("[OK] 图1 inserted")

# ======== 保存 ========
doc.save(OUT_PATH)
print(f"[DONE] 已保存到: {OUT_PATH}")
